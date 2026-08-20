"""臺灣漁業 Hybrid-RAG：清洗、階層切塊、Ollama 向量與混合檢索。

架構重點
--------
- 資料清洗：頁首/頁尾去除、OCR 術語校正、學名格式統一
- 文件擷取：PDF（表格+圖片）、DOCX、DOC（Word COM→DOCX）、ODT（含表格）
- 切塊策略：階層式切塊 300–500 字，50–80 字跨頁 Overlap
- 向量嵌入：語意標籤前綴（[魚種][病害][類型]）+ 純文字，降低 JSON 噪音
- BM25 索引：Unigram + Bigram + Trigram，對長專業術語召回更準
- 混合檢索：語意向量（BGE-M3）+ BM25 RRF 融合，專有名詞時提升 Lexical 權重
- Metadata 後過濾：依魚種、文件類型、病害篩選
"""
from __future__ import annotations

import json
import math
import os
import re
import sqlite3
import tempfile
import zipfile
from array import array
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable
from xml.etree import ElementTree as ET

import fitz
import requests
from docx import Document

# ── 路徑與環境 ────────────────────────────────────────────────────────────────
ROOT = Path(__file__).resolve().parent
DATA_DIR = ROOT / "my_data"
DB_PATH = ROOT / "rag_store.sqlite3"
ASSET_DIR = ROOT / "rag_assets"

OLLAMA_HOST = os.environ.get("OLLAMA_HOST", "http://127.0.0.1:11434").rstrip("/")
EMBED_MODEL = os.environ.get("EMBED_MODEL", "bge-m3:latest")
MODELS = ("gemma2:9b", "deepseek-r1:8b", "qwen2.5-coder:7b", "deepseek-r1:7b")
SUPPORTED = {".pdf", ".doc", ".docx", ".odt", ".txt", ".md"}

CHUNK_TARGET = 420   # 目標切塊字數
CHUNK_OVERLAP = 70   # 跨塊 Overlap 字數
BATCH_SIZE = 24      # 每批嵌入數量

# ── 術語校正清單 ──────────────────────────────────────────────────────────────
# OCR 誤讀、簡繁混用、縮寫 → 標準繁中用語
TERM_CORRECTIONS: dict[str, str] = {
    # 吳郭魚
    "吳鍋魚": "吳郭魚",
    "吳郭漁": "吳郭魚",
    "吴郭鱼": "吳郭魚",
    "吳郭鱼": "吳郭魚",
    # 虱目魚
    "虱目漁": "虱目魚",
    "虱目渔": "虱目魚",
    "虱目仔": "虱目魚",
    # 石斑魚
    "石班魚": "石斑魚",
    "石斑仔": "石斑魚",
    # 鱸魚類
    "金目鱸魚": "金目鱸",
    "七星鱸魚": "七星鱸",
    "海鱸": "七星鱸",
    "花鱸": "七星鱸",
    "鱸仔魚": "金目鱸",
    # 烏魚
    "烏仔魚": "烏魚",
    # 蝦類（學名備註）
    "泰國蝦": "泰國蝦（*Macrobrachium rosenbergii*）",
    "草蝦": "草蝦（*Penaeus monodon*）",
    # 鰻魚
    "鰻苗": "鰻魚苗",
    "鰻仔": "鰻魚",
    # 常見符號誤讀
    "\uff0d": "-",
    "\uff0e": ".",
}

# ── 魚種與病害清單（Metadata 標記用）────────────────────────────────────────
SPECIES = (
    "石斑魚", "吳郭魚", "虱目魚", "鰻魚", "金目鱸",
    "七星鱸", "烏魚", "泰國蝦", "草蝦", "鱸魚",
)
DISEASES = (
    "虹彩病毒", "病毒性神經壞死", "弧菌病", "白點病",
    "寄生蟲病", "黴菌病", "鰓病", "細菌性疾病",
    "愛德華氏菌", "鏈球菌", "水黴病", "白頭白嘴病",
    "沙門氏菌", "立克次體", "諾卡氏菌",
)
DOC_TYPE_RULES: dict[str, tuple[str, ...]] = {
    "用藥法規": ("藥", "殘留", "法規", "規範", "標準", "許可", "禁止", "法條"),
    "病害圖鑑": ("病", "疾病", "寄生蟲", "黴菌", "病毒", "防治", "症狀", "治療", "感染"),
    "養殖手冊": ("養殖", "技術手冊", "管理", "種苗", "飼養", "飼料", "水質", "繁殖"),
}

# ── 章節標題偵測 ──────────────────────────────────────────────────────────────
HEADING = re.compile(
    r"^(?:"
    r"第[一二三四五六七八九十百千\d]+[章節篇條款]"   # 第X章/節/篇
    r"|[一二三四五六七八九十]+[、．.、]"              # 一、二、
    r"|\d+(?:\.\d+){0,3}[\s\u3000]"                  # 1.2.3
    r"|[（(][一二三四五六七八九十\d]+[）)]"           # （一）(1)
    r").{0,80}$"
)

# ── 頁首/頁尾噪音 Pattern ─────────────────────────────────────────────────────
_NOISE: list[re.Pattern] = [
    re.compile(r"行政院農業委員會|農業部水產試驗所|水產試驗所|行政院|農業部"),
    re.compile(r"版權所有|copyright|all rights reserved", re.I),
    re.compile(r"^[\-=─━‐＝]{3,}$"),                  # 水平分隔線
    re.compile(r"^\d{4}\s*年\s*\d{1,2}\s*月$"),        # 純日期行
    re.compile(r"^[第\s]*\d{1,4}\s*[/／]\s*\d{1,4}[頁\s]*$"),  # 頁碼 N/M
    re.compile(r"^[第\s]*\d{1,4}\s*頁$"),              # 第N頁
    re.compile(r"^[─—\-]\s*\d+\s*[─—\-]$"),            # — N —
]


# ── 資料結構 ──────────────────────────────────────────────────────────────────
@dataclass
class PageRecord:
    page: int | None
    text: str
    tables: list[str]
    images: list[str]


@dataclass
class Chunk:
    id: int | None
    source: str
    page: int | None
    position: int
    content: str
    metadata: dict
    score: float = 0.0


# ── 文字工具 ──────────────────────────────────────────────────────────────────

def clean_text(text: str) -> str:
    """術語校正 + 空白正規化（不觸動換行）。"""
    text = (
        text.replace("\u3000", " ")
            .replace("\x00", "")
            .replace("\r\n", "\n")
            .replace("\r", "\n")
    )
    for wrong, right in TERM_CORRECTIONS.items():
        text = text.replace(wrong, right)
    return re.sub(r"[ \t]+", " ", text).strip()


def cjk_tokens(text: str) -> str:
    """Unigram + Bigram + Trigram for CJK；ASCII 保留 2-gram 以上整詞。
    三層 n-gram 確保長專業術語（如「虹彩病毒」）在 FTS5 中有足夠召回率。
    """
    norm = re.sub(r"\s+", "", text.lower())
    tokens: list[str] = list(norm)                                 # unigram
    tokens += [norm[i : i + 2] for i in range(len(norm) - 1)]     # bigram
    tokens += [norm[i : i + 3] for i in range(len(norm) - 2)]     # trigram
    tokens += re.findall(r"[a-z0-9_]{2,}", text.lower())           # ASCII 詞
    return " ".join(tokens)


def markdown_table(rows: list[list[object]]) -> str:
    """將二維陣列轉為 Markdown 表格；過濾全空列、防止 pipe 衝突。"""
    cleaned = [
        [clean_text(str(c or "")).replace("|", "\\|") for c in row]
        for row in rows
        if any(str(c).strip() for c in row)
    ]
    if not cleaned:
        return ""
    width = max(len(r) for r in cleaned)
    cleaned = [r + [""] * (width - len(r)) for r in cleaned]
    header = "| " + " | ".join(cleaned[0]) + " |"
    sep = "| " + " | ".join(["---"] * width) + " |"
    body = "\n".join("| " + " | ".join(r) + " |" for r in cleaned[1:])
    return f"{header}\n{sep}\n{body}" if body else f"{header}\n{sep}"


# ── 頁首/頁尾去除 ────────────────────────────────────────────────────────────

def _is_noise_line(line: str) -> bool:
    s = line.strip()
    if not s or len(s) < 2:
        return True
    # 純頁碼
    if re.fullmatch(r"(?:第\s*)?\d{1,4}", s):
        return True
    return any(pat.search(s) for pat in _NOISE)


def _repeated_lines(records: list[PageRecord]) -> set[str]:
    """偵測跨頁重複行（頻率 ≥ 60%）作為頁首/頁尾候選。"""
    if len(records) < 3:
        return set()
    seen: dict[str, int] = {}
    for rec in records:
        for line in {clean_text(x) for x in rec.text.splitlines() if 2 < len(clean_text(x)) < 70}:
            seen[line] = seen.get(line, 0) + 1
    return {line for line, cnt in seen.items() if cnt / len(records) >= 0.6}


def remove_page_noise(records: list[PageRecord]) -> list[PageRecord]:
    repeated = _repeated_lines(records)
    result: list[PageRecord] = []
    for rec in records:
        lines = [
            line
            for raw in rec.text.splitlines()
            if (line := clean_text(raw))
            and line not in repeated
            and not _is_noise_line(line)
        ]
        result.append(PageRecord(rec.page, "\n".join(lines), rec.tables, rec.images))
    return result


# ── 文件擷取 ──────────────────────────────────────────────────────────────────

def extract_pdf(path: Path) -> Iterable[PageRecord]:
    """PyMuPDF 擷取：文字 + 表格 Markdown + 圖片存檔。"""
    pdf = fitz.open(path)
    asset_base = ASSET_DIR / path.stem
    try:
        for number, page in enumerate(pdf, 1):
            tables: list[str] = []
            images: list[str] = []
            try:
                tables = [t for t in (markdown_table(tbl.extract()) for tbl in page.find_tables().tables) if t]
            except Exception:
                pass
            for idx, img in enumerate(page.get_images(full=True), 1):
                try:
                    extracted = pdf.extract_image(img[0])
                    asset_base.mkdir(parents=True, exist_ok=True)
                    target = asset_base / f"p{number:03d}_{idx}.{extracted['ext']}"
                    if not target.exists():
                        target.write_bytes(extracted["image"])
                    images.append(str(target.relative_to(ROOT)).replace("\\", "/"))
                except Exception:
                    continue
            yield PageRecord(number, page.get_text("text"), tables, images)
    finally:
        pdf.close()


def extract_docx(path: Path) -> Iterable[PageRecord]:
    """python-docx 擷取段落文字與表格。"""
    doc = Document(path)
    tables = [markdown_table([[c.text for c in r.cells] for r in t.rows]) for t in doc.tables]
    text = "\n".join(p.text for p in doc.paragraphs)
    yield PageRecord(None, text, [t for t in tables if t], [])


def _doc_fallback_text(path: Path) -> str:
    """Word 不可用時的備援：從 .doc binary 正則萃取 CJK 與 ASCII 文字片段。
    精度有限，但至少保留關鍵術語供索引使用。
    """
    raw = path.read_bytes()
    # 嘗試 cp950/utf-16-le 解碼，取出可讀字元
    texts: list[str] = []
    for enc in ("utf-16-le", "cp950", "gbk", "latin-1"):
        try:
            decoded = raw.decode(enc, errors="ignore")
            # 擷取連續 CJK 或 ASCII 可讀片段（長度 >= 4）
            chunks = re.findall(r"[\u4e00-\u9fff\u3400-\u4dbf！，。、？；「」『』【】\w]{4,}", decoded)
            if len(chunks) > 10:
                texts = chunks
                break
        except Exception:
            continue
    return "\n".join(texts)


def extract_doc_win32(path: Path) -> Iterable[PageRecord]:
    """舊版 .doc 擷取。
    優先：Word COM 物件另存為 .docx → python-docx（保留表格）。
    備援：Microsoft Word 不可用時，用 binary regex 萃取 CJK 文字片段。
    """
    try:
        import win32com.client  # type: ignore
    except ImportError:
        # pywin32 未安裝，直接走備援
        yield PageRecord(None, _doc_fallback_text(path), [], [])
        return

    try:
        word = win32com.client.DispatchEx("Word.Application")
    except Exception:
        # Word 未安裝或 COM 物件建立失敗，走備援
        yield PageRecord(None, _doc_fallback_text(path), [], [])
        return

    try:
        word.Visible = False
        with tempfile.TemporaryDirectory() as tmp:
            docx_out = Path(tmp) / "converted.docx"
            com_doc = word.Documents.Open(str(path.resolve()), ReadOnly=True)
            try:
                com_doc.SaveAs2(str(docx_out), FileFormat=16)  # 16 = wdFormatDocumentDefault
            finally:
                com_doc.Close(False)
            yield from extract_docx(docx_out)
    except Exception:
        # Word 開啟/儲存失敗，走備援
        yield PageRecord(None, _doc_fallback_text(path), [], [])
    finally:
        try:
            word.Quit()
        except Exception:
            pass


def extract_odt(path: Path) -> Iterable[PageRecord]:
    """ODT (ZIP+XML) 擷取：段落文字 + 表格（避免重複計入 table 內的 text:p）。"""
    _T = "urn:oasis:names:tc:opendocument:xmlns:text:1.0"
    _TBL = "urn:oasis:names:tc:opendocument:xmlns:table:1.0"

    with zipfile.ZipFile(path) as archive:
        root = ET.fromstring(archive.read("content.xml"))

    tables: list[str] = []
    # 記錄所有「屬於表格內部」的 text:p 元素 id，避免重複提取
    table_para_ids: set[int] = set()

    for tbl_elem in root.iter(f"{{{_TBL}}}table"):
        rows: list[list[str]] = []
        for row_elem in tbl_elem.iter(f"{{{_TBL}}}table-row"):
            cells: list[str] = []
            for cell_elem in row_elem.iter(f"{{{_TBL}}}table-cell"):
                cell_text = " ".join(t.strip() for t in cell_elem.itertext() if t.strip())
                cells.append(cell_text)
            if any(cells):
                rows.append(cells)
        md = markdown_table(rows)
        if md:
            tables.append(md)
        for p in tbl_elem.iter(f"{{{_T}}}p"):
            table_para_ids.add(id(p))

    paragraphs: list[str] = []
    for p in root.iter(f"{{{_T}}}p"):
        if id(p) not in table_para_ids:
            line = " ".join(t.strip() for t in p.itertext() if t.strip())
            if line:
                paragraphs.append(line)

    yield PageRecord(None, "\n".join(paragraphs), tables, [])


def extract_file(path: Path) -> Iterable[PageRecord]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        yield from extract_pdf(path)
    elif suffix == ".docx":
        yield from extract_docx(path)
    elif suffix == ".doc":
        yield from extract_doc_win32(path)
    elif suffix == ".odt":
        yield from extract_odt(path)
    else:
        yield PageRecord(None, path.read_text(encoding="utf-8", errors="ignore"), [], [])


# ── Metadata 與 Embed Text ───────────────────────────────────────────────────

def metadata_for(text: str, source: str, page: int | None, hierarchy: str) -> dict:
    fish = [s for s in SPECIES if s in text or s in source] or ["未標記"]
    disease = [d for d in DISEASES if d in text or d in source] or ["未標記"]
    haystack = f"{source} {text}".lower()
    doc_type = next(
        (label for label, kws in DOC_TYPE_RULES.items() if any(kw in haystack for kw in kws)),
        "一般參考資料",
    )
    return {
        "fish_species": fish,
        "disease_name": disease,
        "doc_type": doc_type,
        "source": f"{source}{f'，第 {page} 頁' if page else ''}",
        "hierarchy": hierarchy or "未標記章節",
    }


def build_embed_text(metadata: dict, text: str) -> str:
    """語意標籤前綴 + 純文字送給向量模型。
    讓 BGE-M3 在嵌入時帶有領域上下文，降低 JSON 噪音影響。
    """
    species = "、".join(metadata["fish_species"])
    disease = "、".join(metadata["disease_name"])
    return f"[魚種:{species}][病害:{disease}][類型:{metadata['doc_type']}]\n{text}"


# ── 階層式切塊 ────────────────────────────────────────────────────────────────

def hierarchical_chunks(
    records: list[PageRecord],
    target: int = CHUNK_TARGET,
    overlap: int = CHUNK_OVERLAP,
) -> Iterable[tuple[int | None, str, str]]:
    """依章節階層切塊；跨頁 Overlap 銜接，確保病害描述不中斷。
    Yields: (page, section_hierarchy, text_chunk)
    """
    hierarchy: list[str] = []
    prev_tail: str = ""  # 上一個 block 的尾巴，帶入下個 block 開頭

    for rec in records:
        blocks: list[tuple[str, str]] = []  # (hierarchy_str, raw_text)
        current: list[str] = []

        for line in rec.text.splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            if HEADING.match(stripped):
                if current:
                    blocks.append((" > ".join(hierarchy), "\n".join(current)))
                    current = []
                hierarchy = (hierarchy + [stripped])[-4:]  # 最多保留 4 層
            else:
                current.append(stripped)

        if current:
            blocks.append((" > ".join(hierarchy), "\n".join(current)))

        # 表格 block（附於當前章節下）
        blocks += [(" > ".join(hierarchy), "【表格資料】\n" + tbl) for tbl in rec.tables]

        # 圖片索引 block
        if rec.images:
            img_table = markdown_table(
                [
                    ["圖鑑圖片", "檔案路徑"],
                    *[
                        [f"第 {rec.page} 頁圖 {i + 1}", f"![圖片]({img})"]
                        for i, img in enumerate(rec.images)
                    ],
                ]
            )
            blocks.append((" > ".join(hierarchy), "【圖片圖鑑索引】\n" + img_table))

        for section, block in blocks:
            # 跨頁/跨 block Overlap 銜接
            if prev_tail:
                block = prev_tail.rstrip() + "\n" + block.lstrip()
                prev_tail = ""

            block = clean_text(block)
            if not block:
                continue

            while block:
                if len(block) <= target + 80:
                    yield rec.page, section, block
                    break
                # 優先在句末自然斷點切塊
                cut = max(block.rfind(mark, target - 80, target + 80) for mark in "。！？；;\n")
                cut = (cut + 1) if cut >= target - 80 else target
                yield rec.page, section, block[:cut].strip()
                block = block[max(1, cut - overlap):].lstrip()

        # 儲存本 record 最後一個 block 的尾巴
        if blocks:
            last_text = clean_text(blocks[-1][1])
            prev_tail = last_text[-overlap:] if len(last_text) > overlap else last_text


# ── SQLite 向量庫 ─────────────────────────────────────────────────────────────

def connect() -> sqlite3.Connection:
    db = sqlite3.connect(DB_PATH)
    db.row_factory = sqlite3.Row
    return db


def setup_db(db: sqlite3.Connection, rebuild: bool) -> None:
    existing = {row[1] for row in db.execute("PRAGMA table_info(chunks)")}
    # 自動升級：若舊版缺少 embedding 欄位則強制重建
    if existing and "embedding" not in existing:
        rebuild = True
    if rebuild:
        db.execute("DROP TABLE IF EXISTS chunks")
        db.execute("DROP TABLE IF EXISTS chunks_fts")
    db.execute(
        """CREATE TABLE IF NOT EXISTS chunks (
            id            INTEGER PRIMARY KEY,
            content       TEXT    NOT NULL,
            source        TEXT    NOT NULL,
            page          INTEGER,
            position      INTEGER NOT NULL,
            hierarchy     TEXT,
            metadata_json TEXT    NOT NULL,
            embedding     BLOB    NOT NULL
        )"""
    )
    db.execute(
        "CREATE VIRTUAL TABLE IF NOT EXISTS chunks_fts "
        "USING fts5(search_text, chunk_id UNINDEXED)"
    )
    db.commit()


def pack_vector(values: list[float]) -> bytes:
    return array("f", values).tobytes()


def unpack_vector(raw: bytes) -> array:
    v: array = array("f")
    v.frombytes(raw)
    return v


# ── Ollama API ────────────────────────────────────────────────────────────────

def check_ollama() -> None:
    """確認 Ollama 服務可達，且嵌入模型已安裝。"""
    try:
        resp = requests.get(f"{OLLAMA_HOST}/api/tags", timeout=10)
        resp.raise_for_status()
    except Exception as exc:
        raise RuntimeError(
            f"無法連線 Ollama（{OLLAMA_HOST}）：{exc}\n"
            "請確認 Ollama 已啟動，或設定 $env:OLLAMA_HOST"
        ) from exc
    installed = [m["name"] for m in resp.json().get("models", [])]
    base = EMBED_MODEL.split(":")[0]
    if not any(base in m for m in installed):
        raise RuntimeError(
            f"找不到向量模型 {EMBED_MODEL}。\n"
            f"請先執行：ollama pull {EMBED_MODEL}\n"
            f"已安裝模型：{', '.join(installed) or '（無）'}"
        )


def embed_texts(texts: list[str]) -> list[list[float]]:
    try:
        resp = requests.post(
            f"{OLLAMA_HOST}/api/embed",
            json={"model": EMBED_MODEL, "input": texts},
            timeout=600,
        )
        resp.raise_for_status()
        return resp.json()["embeddings"]
    except requests.HTTPError as exc:
        detail = exc.response.text[:300] if exc.response is not None else ""
        raise RuntimeError(
            f"向量模型呼叫失敗（{EMBED_MODEL}）：{detail}"
        ) from exc


# ── 建立索引 ──────────────────────────────────────────────────────────────────

def build_index(rebuild: bool = False) -> tuple[int, list[str]]:
    """掃描 my_data/，清洗→切塊→嵌入→寫入 SQLite。"""
    if not DATA_DIR.exists():
        raise FileNotFoundError(f"找不到資料目錄：{DATA_DIR}")

    check_ollama()
    db = connect()
    setup_db(db, rebuild)

    if not rebuild and db.execute("SELECT count(*) FROM chunks").fetchone()[0]:
        raise RuntimeError("索引已存在；重建請加上 --rebuild")

    # (embed_text, content, source, page, position, hierarchy, metadata)
    pending: list[tuple[str, str, str, int | None, int, str, dict]] = []
    errors: list[str] = []

    files = sorted(p for p in DATA_DIR.rglob("*") if p.is_file() and p.suffix.lower() in SUPPORTED)
    print(f"找到 {len(files)} 個文件，開始處理…\n")

    for path in files:
        source = str(path.relative_to(ROOT))
        position = 0
        print(f"  [+] {source}")
        try:
            records = remove_page_noise(list(extract_file(path)))
            for page, hierarchy, text in hierarchical_chunks(records):
                metadata = metadata_for(text, source, page, hierarchy)
                embed_text = build_embed_text(metadata, text)
                pending.append((embed_text, text, source, page, position, hierarchy, metadata))
                position += 1
        except Exception as exc:
            errors.append(f"{source}: {exc}")
            print(f"    [X] ERROR: {exc}")

    print(f"\n共 {len(pending)} 個切塊，開始嵌入（batch={BATCH_SIZE}）…\n")

    for start in range(0, len(pending), BATCH_SIZE):
        batch = pending[start : start + BATCH_SIZE]
        vectors = embed_texts([row[0] for row in batch])
        for row, vec in zip(batch, vectors):
            embed_text, content, source, page, position, hierarchy, metadata = row
            cursor = db.execute(
                "INSERT INTO chunks(content, source, page, position, hierarchy, metadata_json, embedding) "
                "VALUES (?, ?, ?, ?, ?, ?, ?)",
                (
                    content,
                    source,
                    page,
                    position,
                    hierarchy,
                    json.dumps(metadata, ensure_ascii=False),
                    pack_vector(vec),
                ),
            )
            db.execute(
                "INSERT INTO chunks_fts(search_text, chunk_id) VALUES (?, ?)",
                (cjk_tokens(embed_text), cursor.lastrowid),
            )
        db.commit()
        done = min(start + BATCH_SIZE, len(pending))
        print(f"  嵌入進度：{done} / {len(pending)} ({done * 100 // len(pending)}%)")

    db.close()
    return len(pending), errors


# ── 混合檢索 ──────────────────────────────────────────────────────────────────

def cosine(a: array, b: array) -> float:
    dot = sum(x * y for x, y in zip(a, b))
    na = math.sqrt(sum(x * x for x in a))
    nb = math.sqrt(sum(x * x for x in b))
    return dot / (na * nb) if na and nb else 0.0


def retrieve(
    question: str,
    top_k: int = 5,
    fish_species: str | None = None,
    doc_type: str | None = None,
    disease: str | None = None,
) -> list[Chunk]:
    """Hybrid 檢索：BM25 + 語意向量 RRF 融合，支援 Metadata 後過濾。

    Parameters
    ----------
    question    : 使用者問題
    top_k       : 回傳最相關段落數
    fish_species: 限定魚種（None 表示不過濾）
    doc_type    : 限定文件類型（None 表示不過濾）
    disease     : 限定病害名稱（None 表示不過濾）
    """
    db = connect()
    try:
        # ── BM25 Lexical 檢索 ──────────────────────────────────────────────
        fts_query = " OR ".join(cjk_tokens(question).split())
        lexical: dict[int, float] = {
            row["chunk_id"]: 1.0 / (rank + 2)
            for rank, row in enumerate(
                db.execute(
                    "SELECT chunk_id FROM chunks_fts "
                    "WHERE chunks_fts MATCH ? ORDER BY bm25(chunks_fts) LIMIT ?",
                    (fts_query, top_k * 8),
                ),
                1,
            )
        }

        # ── 語意向量檢索 ───────────────────────────────────────────────────
        qvec = array("f", embed_texts([question])[0])
        semantic = sorted(
            (
                (row["id"], cosine(qvec, unpack_vector(row["embedding"])))
                for row in db.execute("SELECT id, embedding FROM chunks")
            ),
            key=lambda x: x[1],
            reverse=True,
        )[: top_k * 8]

        # ── 動態權重：查詢含專有名詞時提升 BM25 權重 ──────────────────────
        has_term = any(t in question for t in list(SPECIES) + list(DISEASES))
        w_sem, w_lex = (0.42, 0.58) if has_term else (0.55, 0.45)

        scores: dict[int, float] = {cid: w_sem * sc for cid, sc in semantic}
        for cid, sc in lexical.items():
            scores[cid] = scores.get(cid, 0.0) + w_lex * sc

        if not scores:
            return []

        ids = tuple(scores)
        rows_map = {
            row["id"]: row
            for row in db.execute(
                f"SELECT * FROM chunks WHERE id IN ({','.join('?' * len(ids))})", ids
            )
        }

        hits: list[Chunk] = []
        for cid, score in scores.items():
            row = rows_map.get(cid)
            if not row:
                continue
            meta = json.loads(row["metadata_json"])

            # ── Metadata 後過濾 ────────────────────────────────────────────
            if fish_species and fish_species not in meta["fish_species"]:
                continue
            if doc_type and doc_type != meta["doc_type"]:
                continue
            if disease and disease not in meta["disease_name"]:
                continue

            # 精準術語 bonus（+0.08）
            if any(t in question for t in meta["fish_species"] + meta["disease_name"]):
                score += 0.08

            hits.append(
                Chunk(row["id"], row["source"], row["page"], row["position"], row["content"], meta, score)
            )

        return sorted(hits, key=lambda c: c.score, reverse=True)[:top_k]
    finally:
        db.close()


def ask_ollama(question: str, model: str, chunks: list[Chunk]) -> str:
    if model not in MODELS:
        raise ValueError(f"不支援的模型：{model}")
    context = "\n\n".join(
        f"[來源 {i + 1}: {chunk.metadata['source']}]\n{chunk.content}"
        for i, chunk in enumerate(chunks)
    )
    prompt = (
        "你是臺灣漁業與水產養殖資料助理，只依據以下參考資料，以繁體中文回答問題。"
        "資料不足時請明確說明無法確認。每項事實以 [來源編號] 標注出處。\n\n"
        f"參考資料：\n{context or '（沒有找到相關資料）'}\n\n"
        f"問題：{question}"
    )
    resp = requests.post(
        f"{OLLAMA_HOST}/api/chat",
        json={
            "model": model,
            "stream": False,
            "messages": [{"role": "user", "content": prompt}],
        },
        timeout=600,
    )
    resp.raise_for_status()
    return resp.json()["message"]["content"]


# ── 知識庫統計 ────────────────────────────────────────────────────────────────

def get_db_stats() -> dict:
    """回傳知識庫摘要統計，供 Streamlit UI 使用。"""
    db = connect()
    try:
        total = db.execute("SELECT count(*) FROM chunks").fetchone()[0]
        if total == 0:
            return {"total_chunks": 0, "total_sources": 0, "doc_types": {}, "species": {}}
        sources = db.execute("SELECT count(DISTINCT source) FROM chunks").fetchone()[0]
        doc_types = dict(
            db.execute(
                "SELECT json_extract(metadata_json, '$.doc_type'), count(*) "
                "FROM chunks GROUP BY json_extract(metadata_json, '$.doc_type')"
            ).fetchall()
        )
        # 魚種分布（metadata 中 fish_species 為 JSON array，用 LIKE 近似統計）
        species_counts: dict[str, int] = {}
        for sp in SPECIES:
            cnt = db.execute(
                "SELECT count(*) FROM chunks WHERE metadata_json LIKE ?",
                (f'%"{sp}"%',),
            ).fetchone()[0]
            if cnt:
                species_counts[sp] = cnt
        return {
            "total_chunks": total,
            "total_sources": sources,
            "doc_types": doc_types,
            "species": species_counts,
        }
    except Exception:
        return {"total_chunks": 0, "total_sources": 0, "doc_types": {}, "species": {}}
    finally:
        db.close()
